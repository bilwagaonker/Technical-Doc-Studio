from pathlib import Path
from datetime import datetime
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class ExportService:

    def __init__(self):
        pass

    def _safe(self, value, default=""):

        if value is None:
            return default

        return str(value)

    def _safe_filename(self, name):

        name = self._safe(name, "SAP_Documentation")

        name = re.sub(r'[\\/*?:"<>|]', "", name)

        name = name.replace(" ", "_")

        return name

    def export_docx(

        self,

        documentation,

        frame_folder,

        output_folder="app/storage/output"

    ):

        output_folder = Path(output_folder)

        output_folder.mkdir(
            parents=True,
            exist_ok=True
        )

        frame_folder = Path(frame_folder)

        document = Document()

        ###############################################################
        # Cover
        ###############################################################

        title = document.add_heading(
            "SAP Documentation",
            level=1
        )

        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = document.add_paragraph()

        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle.add_run(
            documentation.get(
                "title",
                "Technical Process Documentation"
            )
        ).bold = True

        document.add_paragraph(
            f"Generated : {datetime.now():%d-%b-%Y %H:%M}"
        )

        document.add_page_break()

        ###############################################################
        # Purpose
        ###############################################################

        document.add_heading(
            "Purpose",
            level=1
        )

        document.add_paragraph(
            documentation.get(
                "purpose",
                "Purpose not generated."
            )
        )

        ###############################################################
        # Prerequisites
        ###############################################################

        prereq = documentation.get(
            "prerequisites",
            []
        )

        if prereq:

            document.add_heading(
                "Prerequisites",
                level=1
            )

            for item in prereq:

                document.add_paragraph(
                    str(item),
                    style="List Bullet"
                )

        ###############################################################
        # Procedure
        ###############################################################

        document.add_heading(
            "Procedure",
            level=1
        )

        steps = documentation.get(
            "steps",
            []
        )

        if not steps:

            document.add_paragraph(
                "No procedural steps generated."
            )

        for index, step in enumerate(steps, start=1):

            document.add_heading(
                f"Step {step.get('step', index)}",
                level=2
            )

            if step.get("title"):

                run = document.add_paragraph().add_run(
                    step["title"]
                )

                run.bold = True

            document.add_paragraph(
                step.get(
                    "description",
                    ""
                )
            )

            ###########################################################
            # SAP Fields
            ###########################################################

            fields = step.get(
                "fields",
                []
            )

            if fields:

                document.add_heading(
                    "Fields",
                    level=3
                )

                table = document.add_table(
                    rows=1,
                    cols=2
                )

                table.style = "Light Grid"

                hdr = table.rows[0].cells

                hdr[0].text = "Field"

                hdr[1].text = "Description"

                for field in fields:

                    row = table.add_row().cells

                    if isinstance(field, dict):

                        row[0].text = self._safe(
                            field.get("name")
                        )

                        row[1].text = self._safe(
                            field.get("description")
                        )

                    else:

                        row[0].text = self._safe(field)

                        row[1].text = ""

            ###########################################################
            # Expected Result
            ###########################################################

            if step.get("expected_result"):

                document.add_heading(
                    "Expected Result",
                    level=3
                )

                document.add_paragraph(
                    step["expected_result"]
                )

            ###########################################################
            # Screenshot
            ###########################################################

            image_name = step.get("image")

            if image_name:

                image_path = frame_folder / image_name
                print("=" * 60)
                print("Image from documentation :", image_name)
                print("Frame folder :", frame_folder)

                if image_path.exists():

                    print("ADDING IMAGE")
                    document.add_picture(
                        str(image_path),
                        width=Inches(6.3)
                    )
                else:
                    print("PIC NOT FOUND")

                    caption = document.add_paragraph(
                        f"Figure {index}"
                    )

                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        ###############################################################
        # Notes
        ###############################################################

        notes = documentation.get(
            "notes",
            []
        )

        if notes:

            document.add_heading(
                "Notes",
                level=1
            )

            for note in notes:

                document.add_paragraph(
                    str(note),
                    style="List Bullet"
                )

        ###############################################################
        # Warnings
        ###############################################################

        warnings = documentation.get(
            "warnings",
            []
        )

        if warnings:

            document.add_heading(
                "Warnings",
                level=1
            )

            for warning in warnings:

                document.add_paragraph(
                    str(warning),
                    style="List Bullet"
                )

        ###############################################################
        # Footer
        ###############################################################

        section = document.sections[0]

        footer = section.footer

        footer.paragraphs[0].text = (
            "Generated by AI Technical Documentation Studio"
        )

        ###############################################################
        # Save
        ###############################################################

        filename = self._safe_filename("SAP Documentation")

        output_file = output_folder / f"{filename}.docx"

        document.save(output_file)

        return output_file


    def export_markdown(self, documentation,output_folder="app/storage/output"):

        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)

        filename = self._safe_filename("SAP Documentation")

        output_file = output_folder / f"{filename}.md"

        md = []

        md.append(f"# {"SAP Documentation"}\n")

        md.append(f"## {documentation.get('title','')}\n")

        md.append("## Purpose\n")
        md.append(documentation.get("purpose", "") + "\n")

        if documentation.get("prerequisites"):

            md.append("## Prerequisites\n")

            for item in documentation["prerequisites"]:
                md.append(f"- {item}")

            md.append("")

        md.append("## Procedure\n")

        for step in documentation.get("steps", []):

            md.append(f"### Step {step.get('step')}")

            if step.get("title"):
                md.append(f"**{step['title']}**")

            md.append(step.get("description", ""))

            if step.get("image"):
                md.append(
                f"![Screenshot](../frames/{step['image']})"
                )

            md.append("")

        if documentation.get("notes"):

            md.append("## Notes")

            for note in documentation["notes"]:
                md.append(f"- {note}")

        if documentation.get("warnings"):

            md.append("\n## Warnings")

            for warning in documentation["warnings"]:
                md.append(f"- {warning}")

        output_file.write_text(
            "\n".join(md),
            encoding="utf8"
        )

        return output_file
    
    def export_html(self,documentation,output_folder="app/storage/output"):

        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)

        filename = self._safe_filename("SAP Documentation")

        output_file = output_folder / f"{filename}.html"

        html = []

        html.append("""
<html>
<head>
<title>SAP Documentation</title>

<style>

body{
font-family:Arial;
margin:40px;
line-height:1.6;
}

img{
width:900px;
border:1px solid #ccc;
margin-top:10px;
margin-bottom:20px;
}

.step{
margin-bottom:40px;
}

</style>

</head>
<body>
""")

        html.append(f"<h1>{"SAP Documentation"}</h1>")

        html.append(f"<h2>{documentation.get('title','')}</h2>")

        html.append("<h3>Purpose</h3>")

        html.append(
        f"<p>{documentation.get('purpose','')}</p>"
        )

        html.append("<h2>Procedure</h2>")

        for step in documentation.get("steps", []):

            html.append("<div class='step'>")

            html.append(
            f"<h3>Step {step.get('step')}</h3>"
            )

            html.append(
                f"<b>{step.get('title','')}</b>"
            )

            html.append(
                f"<p>{step.get('description','')}</p>"
            )

            if step.get("image"):

                html.append(
                    f"<img src='../frames/{step['image']}'>"
                )

            html.append("</div>")

        html.append("</body></html>")

        output_file.write_text(
        "\n".join(html),
        encoding="utf8"
        )

        return output_file
    
    